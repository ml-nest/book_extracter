import os
import re
import io
import argparse
from urllib.parse import urlparse, urljoin

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import fitz  # PyMuPDF


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _download_image(src: str, image_dir: str, headers: dict, base_url: str,
                    counter: list) -> str | None:
    """Resolve, download and save an image; return its local filepath or None."""
    if src.startswith("//"):
        src = "https:" + src




          
    elif src.startswith("/"):
        parsed = urlparse(base_url)
        src = f"{parsed.scheme}://{parsed.netloc}{src}"
    elif not src.startswith("http"):
        src = urljoin(base_url, src)

    try:
        r = requests.get(src, headers=headers, timeout=15)
        r.raise_for_status()
        filename = os.path.basename(src.split("?")[0]) or "image"
        if not re.search(r"\.\w{2,5}$", filename):
            filename += ".jpg"
        # Avoid filename collisions
        counter[0] += 1
        name, ext = os.path.splitext(filename)
        filepath = os.path.join(image_dir, f"{name}_{counter[0]}{ext}")
        with open(filepath, "wb") as f:
            f.write(r.content)
        print(f"Downloaded: {filepath}")
        return filepath
    except Exception as e:
        print(f"Failed to download {src}: {e}")
        return None


def _image_to_png_buf(img_path: str) -> io.BytesIO:
    """Convert any image (including SVG) to an in-memory PNG buffer."""
    buf = io.BytesIO()
    if img_path.lower().endswith(".svg"):
        doc = fitz.open(img_path)
        pix = doc[0].get_pixmap(matrix=fitz.Matrix(2, 2))
        buf.write(pix.tobytes("png"))
        doc.close()
    else:
        with Image.open(img_path) as img:
            if img.mode in ("P", "RGBA"):
                img = img.convert("RGBA")
                bg = Image.new("RGB", img.size, (255, 255, 255))
                bg.paste(img, mask=img.split()[3])
                img = bg
            elif img.mode != "RGB":
                img = img.convert("RGB")
            img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Extraction — returns a structured list of elements
# ---------------------------------------------------------------------------

# Patterns that indicate junk / MSO style blobs / CSS fragments
_JUNK_RE = re.compile(
    r"(mso-|MsoNormal|font-family|mso-style|@page|table\.Mso"
    r"|mso-padding|widow-orphan|line-height:\s*\d"
    r"|\{mso-|font-size:\s*\d|margin-bottom:\s*\d"
    r"|Rs\.\s*\d|View\s*&\s*Get|Year\s*Validity"
    r"|Price\s*Per\s*Month|Offline\s*Support|3/6/12-Month"
    r"|EN-US|X-NONE|Style\s+Definitions"
    r"|ICSE ML Aggarwal Solutions Class 7 Mathematics"
    r"|NEET.*?Coaching|Coaching\s+Programs"
    r"|\U0001F4CA|\U0001F4F9|\U0001F393|\U0001F3AF"
    r"|MCQ.*?PYQ|Study\s+Material.*?Syllabus"
    r"|National\s+Eligibility.*?Entrance"
    r"|mso-ascii|mso-hansi|minor-latin"
    r"|mso-tstyle|mso-style-priority)",
    re.IGNORECASE | re.DOTALL
)


# Prefixes that are MSO style-name artifacts prepended to real content
_MSO_PREFIX_RE = re.compile(
    r"^\s*(Normal|Heading\s*\d*|Default|Body\s*Text)\s*\*{0,2}\s*",
    re.IGNORECASE
)
# Trailing/leading bold markers that sometimes leak in
_BOLD_MARKER_RE = re.compile(r"^\*{1,2}|\*{1,2}$")


def _clean_text(text: str) -> str:
    """Strip MSO style-name prefixes and stray markdown bold markers."""
    text = _MSO_PREFIX_RE.sub("", text)
    text = _BOLD_MARKER_RE.sub("", text).strip()
    return text


def _is_junk(text: str) -> bool:
    """Return True if the text looks like CSS/MSO metadata or promo noise."""
    if not text:
        return True
    # Very short tokens that are just numbers or isolated punctuation
    if len(text) < 3:
        return True
    # Matches known junk patterns
    if _JUNK_RE.search(text):
        return True
    # Lines that are entirely digits / whitespace / punctuation
    stripped = text.strip()
    if re.fullmatch(r'[\d\s.,;:!?()\[\]{}\'"\-_/\\@#$%^&*+=<>|~`]+', stripped):
        return True
    # MSO XML tokens: standalone 'false', 'true', language codes like 'JA'
    if re.fullmatch(r'(false|true|Normal|\w{2}-\w+)', stripped, re.IGNORECASE):
        return True
    # Suspiciously high ratio of whitespace/newlines (MSO blobs have lots of spaces)
    if len(stripped) > 0 and (text.count('\n') / max(len(text), 1)) > 0.3:
        return True
    return False


def _walk(tag, elements: list, image_dir: str, headers: dict,
          base_url: str, counter: list) -> None:
    """Recursively walk HTML and emit structured elements."""
    SKIP = {"script", "style", "nav", "footer", "header", "noscript",
            "aside", "form", "button", "meta", "link"}
    HEADING = {"h1", "h2", "h3", "h4", "h5", "h6"}

    for child in tag.children:
        name = getattr(child, "name", None)
        if name is None:           # NavigableString — skip bare text nodes here
            continue
        if name in SKIP:
            continue

        if name in HEADING:
            text = _clean_text(child.get_text(separator=" ", strip=True))
            if text and not _is_junk(text):
                elements.append({"type": name, "text": text})

        elif name == "p":
            text = _clean_text(child.get_text(separator=" ", strip=True))
            if text and not _is_junk(text):
                elements.append({"type": "p", "text": text})

        elif name in ("ul", "ol"):
            ordered = name == "ol"
            for li in child.find_all("li", recursive=False):
                text = _clean_text(li.get_text(separator=" ", strip=True))
                if text and not _is_junk(text):
                    elements.append({"type": "li", "text": text, "ordered": ordered})

        elif name == "img":
            src = child.get("src") or child.get("data-src")
            if src:
                path = _download_image(src, image_dir, headers, base_url, counter)
                if path:
                    elements.append({"type": "img", "path": path})

        elif name == "table":
            # Render table rows as tab-separated text paragraphs
            for row in child.find_all("tr"):
                cells = [td.get_text(separator=" ", strip=True)
                         for td in row.find_all(["td", "th"])]
                if any(cells):
                    elements.append({"type": "table_row",
                                     "cells": cells,
                                     "is_header": bool(row.find("th"))})
        else:
            # Recurse into div, section, article, main, span, etc.
            _walk(child, elements, image_dir, headers, base_url, counter)


def extract_content_from_url(url: str, image_dir: str = "images") -> list:
    """Fetch a webpage and return a structured list of content elements."""
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        )
    }
    response = requests.get(url, headers=headers, timeout=15)

    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")
    # Remove clearly non-content tags
    for tag in soup(["script", "style", "nav", "footer", "header",
                     "noscript", "aside", "iframe", "form"]):
        tag.decompose()

    # Remove MSO / Word style blobs (appear as <p> or <div> with MSO content)
    # Collect first, then decompose — avoids NoneType on decomposed children
    to_remove = []
    for tag in soup.find_all(True):
        if tag is None or not hasattr(tag, "attrs") or tag.attrs is None:
            continue
        cls = " ".join(tag.get("class") or [])
        if re.search(r"mso|MsoNormal|WordSection", cls):
            to_remove.append(tag)
            continue
        if tag.name == "p":
            raw = tag.get_text()
            if _is_junk(raw):
                to_remove.append(tag)
    for tag in to_remove:
        try:
            tag.decompose()
        except Exception:
            pass

    os.makedirs(image_dir, exist_ok=True)
    elements: list = []
    counter = [0]

    root = (
        soup.find("main")
        or soup.find("article")
        or soup.find(id=re.compile(r"content|post|article|entry", re.I))
        or soup.find(class_=re.compile(
            r"post-content|entry-content|article-body|main-content"
            r"|post_content|article_content|page-content", re.I))
        or soup.find("body")
        or soup
    )
    _walk(root, elements, image_dir, headers, url, counter)

    return elements


# ---------------------------------------------------------------------------
# Word export
# ---------------------------------------------------------------------------

# ── Colour palette ──────────────────────────────────────────────────────────
_C_H1      = RGBColor(0x1F, 0x39, 0x64)   # deep navy
_C_H2      = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
_C_H3      = RGBColor(0x2E, 0x74, 0xB5)   # same blue, smaller
_C_BODY    = RGBColor(0x26, 0x26, 0x26)   # near-black
_C_CAPTION = RGBColor(0x75, 0x75, 0x75)   # grey
_C_TH_BG   = "1F3964"                     # header row fill (hex, no #)
_C_TR_BG   = "DEEAF1"                     # alternate row fill
_C_BORDER  = "2E74B5"                     # table border colour
_FONT      = "Calibri"


def _set_cell_bg(cell, hex_color: str) -> None:
    """Fill a table cell background with a solid colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_table_borders(table, hex_color: str = "2E74B5") -> None:
    """Apply uniform borders to a table."""
    tbl    = table._tbl
    tblPr  = tbl.tblPr
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        tblBdr.append(el)
    tblPr.append(tblBdr)


def _add_h_rule(doc: Document) -> None:
    """Insert a thin blue horizontal rule paragraph."""
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), _C_BORDER)
    pBdr.append(bot)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(4)


def save_to_word(elements: list, output_path: str = "output.docx") -> None:
    """Save structured content elements to a formatted Word document."""
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Track consecutive table rows to build a real docx table
    _pending_rows: list[dict] = []

    def _flush_table() -> None:
        """Convert pending table rows into a real Word table."""
        if not _pending_rows:
            return
        cols = max(len(r["cells"]) for r in _pending_rows)
        tbl  = doc.add_table(rows=0, cols=cols)
        tbl.style = "Table Grid"
        _set_table_borders(tbl, _C_BORDER)

        for i, row_data in enumerate(_pending_rows):
            row   = tbl.add_row()
            cells = row_data["cells"]
            is_hdr = row_data.get("is_header", False)
            for j, text in enumerate(cells):
                cell = row.cells[j]
                cell.text = text
                run  = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(text)
                run.font.name = _FONT
                run.font.size = Pt(10)
                if is_hdr:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_cell_bg(cell, _C_TH_BG)
                elif i % 2 == 0:
                    _set_cell_bg(cell, _C_TR_BG)
        _pending_rows.clear()
        doc.add_paragraph()  # breathing space after table

    HEADING_LEVEL = {"h1": 0, "h2": 1, "h3": 2, "h4": 3, "h5": 3, "h6": 3}

    for el in elements:
        etype = el["type"]

        # Flush any buffered table rows when a non-table element arrives
        if etype != "table_row":
            _flush_table()

        # ── Headings ────────────────────────────────────────────────────────
        if etype in HEADING_LEVEL:
            level = HEADING_LEVEL[etype]
            para  = doc.add_heading("", level=level)
            run   = para.add_run(el["text"])
            run.font.name = _FONT
            if level == 0:
                run.font.size  = Pt(20)
                run.font.color.rgb = _C_H1
                run.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(16)
                para.paragraph_format.space_after  = Pt(4)
                _add_h_rule(doc)
            elif level == 1:
                run.font.size  = Pt(15)
                run.font.color.rgb = _C_H2
                run.bold = True
                para.paragraph_format.space_before = Pt(14)
                para.paragraph_format.space_after  = Pt(4)
            else:
                run.font.size  = Pt(12)
                run.font.color.rgb = _C_H3
                run.bold = True
                para.paragraph_format.space_before = Pt(10)
                para.paragraph_format.space_after  = Pt(3)

        # ── Body paragraph ──────────────────────────────────────────────────
        elif etype == "p":
            para = doc.add_paragraph()
            para.style = doc.styles["Normal"]
            para.paragraph_format.space_after        = Pt(7)
            para.paragraph_format.line_spacing       = Pt(14)
            para.paragraph_format.first_line_indent  = Pt(0)
            run = para.add_run(el["text"])
            run.font.size  = Pt(11)
            run.font.name  = _FONT
            run.font.color.rgb = _C_BODY

        # ── Lists ────────────────────────────────────────────────────────────
        elif etype == "li":
            style = "List Number" if el.get("ordered") else "List Bullet"
            para  = doc.add_paragraph(el["text"], style=style)
            para.paragraph_format.space_after        = Pt(3)
            para.paragraph_format.left_indent        = Inches(0.3)
            for run in para.runs:
                run.font.size  = Pt(11)
                run.font.name  = _FONT
                run.font.color.rgb = _C_BODY

        # ── Table row (buffered) ─────────────────────────────────────────────
        elif etype == "table_row":
            _pending_rows.append(el)

        # ── Image ─────────────────────────────────────────────────────────────
        elif etype == "img":
            img_path = el["path"]
            try:
                buf  = _image_to_png_buf(img_path)

                # Determine dimensions to cap height at 2.32"
                buf.seek(0)
                with Image.open(buf) as _img:
                    img_w_px, img_h_px = _img.size
                buf.seek(0)

                max_width  = Inches(5.0)
                max_height = Inches(2.32)

                if img_w_px > 0 and img_h_px > 0:
                    aspect = img_w_px / img_h_px
                    # Start with full width
                    width  = max_width
                    height = width / aspect
                    # If height exceeds cap, scale down to fit height
                    if height > max_height:
                        height = max_height
                        width  = height * aspect
                else:
                    width, height = max_width, max_height

                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(10)
                para.paragraph_format.space_after  = Pt(4)
                run  = para.add_run()
                run.add_picture(buf, width=int(width), height=int(height))

                cap     = doc.add_paragraph(os.path.basename(img_path))
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(12)
                cap_run = cap.runs[0]
                cap_run.italic         = True
                cap_run.font.size      = Pt(9)
                cap_run.font.name      = _FONT
                cap_run.font.color.rgb = _C_CAPTION
                print(f"Added image: {img_path}")
            except Exception as e:
                print(f"Could not add image {img_path}: {e}")

    _flush_table()  # flush any trailing table rows

    doc.save(output_path)
    print(f"\nWord document saved to: {output_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Extract book content from a base URL pattern and part count."
    )
    parser.add_argument(
        "--last-link",
        dest="last_link",
        help="Last part URL. The script will derive BASE URL, PARTS, and default output name from it.",
    )
    parser.add_argument(
        "--base-url",
        dest="base_url",
        help="Base URL pattern. Use '{}' where the part number should be inserted.",
    )
    parser.add_argument(
        "--parts",
        dest="parts",
        type=int,
        help="Number of parts to fetch.",
    )
    parser.add_argument(
        "--output",
        dest="output",
        default="output.docx",
        help="Output Word document path.",
    )
    return parser.parse_args()


def _derive_inputs_from_last_link(last_link: str) -> tuple[str, int, str]:
    match = re.search(r"^(.*-Part-)(\d+)(\.html?)$", last_link.strip(), re.IGNORECASE)
    if not match:
        raise ValueError(
            "Last link must end with a part number like '-Part-16.html'."
        )

    base_url = f"{match.group(1)}{{}}{match.group(3)}"
    parts = int(match.group(2))

    chapter_match = re.search(
        r"(Chapter-\d+-[^/]+)-Part-\d+\.html?$",
        last_link.strip(),
        re.IGNORECASE,
    )
    if chapter_match:
        output_path = f"{chapter_match.group(1)}.docx"
    else:
        output_path = "output.docx"

    return base_url, parts, output_path


def _resolve_inputs() -> tuple[str, int, str]:
    args = _parse_args()

    if args.last_link:
        base_url, parts, derived_output = _derive_inputs_from_last_link(args.last_link)
        output_path = args.output if args.output != "output.docx" else derived_output
        return base_url, parts, output_path

    last_link = input("Enter last part link (or press Enter to use BASE URL): ").strip()
    if last_link:
        base_url, parts, derived_output = _derive_inputs_from_last_link(last_link)
        output_path = args.output if args.output != "output.docx" else derived_output
        return base_url, parts, output_path

    base_url = args.base_url or input("Enter BASE URL pattern: ").strip()
    while "{}" not in base_url:
        print("BASE URL must include '{}' for the part number.")
        base_url = input("Enter BASE URL pattern: ").strip()

    parts = args.parts
    while parts is None or parts < 1:
        raw_parts = input("Enter PARTS count: ").strip()
        try:
            parts = int(raw_parts)
        except ValueError:
            parts = None
        if parts is None or parts < 1:
            print("PARTS must be a positive integer.")

    return base_url, parts, args.output


if __name__ == "__main__":
    BASE_URL, PARTS, OUTPUT_PATH = _resolve_inputs()

    all_elements: list = []
    for part in range(1, PARTS + 1):
        url        = BASE_URL.format(part)
        image_dir  = f"images/part{part}"
        print(f"\n{'='*60}")
        print(f"Fetching Part {part}: {url}")
        print('='*60)
        try:
            elements = extract_content_from_url(url, image_dir=image_dir)
            print(f"  -> {len(elements)} elements extracted")
            # Insert a part-separator heading between parts
            if all_elements and elements:
                all_elements.append({"type": "h1", "text": f"Part {part}"})
            all_elements.extend(elements)
        except Exception as e:
            print(f"  -> Failed to fetch Part {part}: {e}")

    print(f"\nTotal elements across all parts: {len(all_elements)}")
    save_to_word(all_elements, output_path=OUTPUT_PATH)
